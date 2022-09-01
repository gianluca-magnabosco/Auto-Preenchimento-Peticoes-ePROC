import os
import time
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt
from classes.processos import Processos
import win32gui
import win32con
from functions.aux_functions import connectDataBase
from datetime import date


class Documentos(Processos):

    def __init__(self, index):
        self.index = index


    def getDataFromDatabase(self):
        con = connectDataBase()
        cur = con.cursor()
        
        cur.execute(f"SELECT * FROM processos WHERE num_processo = '{self.numero_processos[self.index]}';")

        result = cur.fetchall()

        self.num_processo_final = result[0][0]
        self.cliente_final = result[0][1]
        self.adversa_final = result[0][2]
        self.cidade_final = result[0][3]
        
        diaDeHoje = date.today().strftime("%d/%m/%Y")

        self.dia = diaDeHoje[:2]
        self.mes = diaDeHoje[3:5]
        self.ano = diaDeHoje[8:]

        mesesExtenso = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
        self.mes = int(self.mes)
        self.mes = mesesExtenso[self.mes - 1]


    def addNewLines(self, paragraph, amount):
        cursor = paragraph.add_run()
        for _ in range(amount):
            cursor.add_break()


    def addCenteredText(self, text, bold):
        paragraph = self.document.add_paragraph()
        paragraph.add_run(text).bold = bold
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return paragraph


    def addFooters(self):
        paragraph = self.addCenteredText("Nestes termos,", bold = False)
        paragraph = self.addCenteredText("Pede deferimento.", bold = False)
        paragraph = self.addCenteredText(f"Mafra, {self.dia} de {self.mes} de 20{self.ano}.", bold = False)

        self.addNewLines(paragraph, amount = 2)

        paragraph = self.addCenteredText("MÁRCIO MAGNABOSCO DA SILVA", bold = True)
        paragraph = self.addCenteredText("OAB/SC 9.738 – OAB/PR 20.962", bold = True)

        self.addNewLines(paragraph, amount = 2)

        paragraph = self.addCenteredText("ALINE REWAY RUTHES", bold = True)
        paragraph = self.addCenteredText("OAB/SC 52.034", bold = True)


    def setStyle(self):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        for paragraph in self.document.paragraphs:
            paragraph.style = self.document.styles['Normal']
            paragraph.paragraph_format.space_after = Pt(0)

            if 'Nestes' in paragraph.text:
                paragraph.paragraph_format.space_after = Pt(1.5)
            if 'Pede' in paragraph.text:
                paragraph.paragraph_format.space_after = Pt(1.5)
            if 'Mafra' in paragraph.text:
                paragraph.paragraph_format.space_after = Pt(1.5)
        
        for run in paragraph.runs:
            run.font.size = Pt(12)


    def gerarPeticao(self):

        self.getDataFromDatabase()

        self.document = Document()

        # TITULO
        paragraph = self.addCenteredText(f"EXMO. SR. DR. JUIZ DE DIREITO DA _ª VARA CÍVEL DA COMARCA DE {self.cidade_final.upper()}/SC", bold = True)
        
        self.addNewLines(paragraph, amount = 5)

        # AUTOS
        paragraph = self.document.add_paragraph()
        paragraph.add_run(f"Autos nº: {self.num_processo_final}").bold = True
        paragraph.paragraph_format.line_spacing_rule = 0

        # PARTE ADVERSA
        paragraph = self.document.add_paragraph()
        paragraph.add_run("Parte adversa: ").bold = True
        paragraph.add_run(self.adversa_final.title())

        self.addNewLines(paragraph, amount = 5)

        # CLIENTE
        paragraph = self.document.add_paragraph()
        paragraph.add_run(f"{self.cliente_final}, ").bold = True
        paragraph.add_run("vem respeitosamente à presença de Vossa Excelência através de seu procurador que esta subscreve, expor e requerer o que segue: ")  
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.5)
        paragraphFormat.left_indent = Inches(1.5)
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraphFormat.line_spacing_rule = 0

        self.addNewLines(paragraph, amount = 7)

        # FOOTERS
        self.addFooters()

        # ESTILO
        self.setStyle()

        nome_documento = f"Petição {self.cliente_final}.docx"
        if len(nome_documento) > 42:
            nome_documento = nome_documento[:42] + ".docx"

        path_documento = os.path.join("arquivos/peticoes", nome_documento)
        self.document.save(path_documento)
        os.startfile(path_documento)
        time.sleep(2)
        maximize = win32gui.GetForegroundWindow()
        win32gui.ShowWindow(maximize, win32con.SW_MAXIMIZE)




    def gerarCumprSentenca(self):

        self.getDataFromDatabase()

        self.document = Document()

        # TITULO
        paragraph = self.addCenteredText(f"EXMO. SR. JUIZ FEDERAL DA _ª VARA FEDERAL DE {self.cidade_final.upper()} - SEÇÃO JUDICIÁRIA DE SANTA CATARINA", bold = True)
        
        self.addNewLines(paragraph, amount = 6)

        # AUTOS
        paragraph = self.document.add_paragraph()
        paragraph.add_run(f"Autos nº: {self.num_processo_final}").bold = True
        paragraph.paragraph_format.line_spacing_rule = 0

        # CUMPR SENTENÇA
        paragraph = self.document.add_paragraph()
        paragraph.add_run("Cumprimento de Sentença").bold = True
        paragraph.paragraph_format.line_spacing_rule = 0

        # LINE
        paragraph = self.document.add_paragraph()
        paragraph.add_run("_______________________________").bold = True
        paragraph.paragraph_format.line_spacing_rule = 0

        self.addNewLines(paragraph, amount = 12)

        # CLIENTE
        paragraph = self.document.add_paragraph()
        paragraph.add_run(f"{self.cliente_final}, ").bold = True
        paragraph.add_run("____________, através dos procuradores que a esta subscrevem, MÁRCIO MAGNABOSCO DA SILVA, advogado inscrito na Ordem dos Advogados do Brasil sob nº 9.738/SC e 20.962/PR e ALINE REWAY RUTHES, advogada inscrita na Ordem dos Advogados do Brasil sob nº 52.034/SC, com escritório profissional na rua Felipe Schmidt, nº 354, conjunto nº 01, Mafra/SC, vem, respeitosamente perante Vossa Excelência, nos termos do art. 513 e seguintes do CPC, requerer")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.25)
        paragraphFormat.left_indent = Inches(1.5)
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraphFormat.line_spacing_rule = 0

        self.addNewLines(paragraph, amount = 2)

        # CUMPR SENTENÇA
        paragraph = self.document.add_paragraph()
        paragraph.add_run("CUMPRIMENTO DE SENTENÇA").bold = True
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.left_indent = Inches(1.5)
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.addNewLines(paragraph, amount = 2)

        # EM FACE DE
        paragraph = self.document.add_paragraph("em face de")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.1)
        paragraphFormat.left_indent = Inches(1.5)

        self.addNewLines(paragraph, amount = 2)

        # PARTE ADVERSA
        paragraph = self.document.add_paragraph()
        paragraph.add_run(f"{self.adversa_final}, ").bold = True
        paragraph.add_run("______ em razão dos fatos e fundamentos a seguir aduzidos: ")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.25)
        paragraphFormat.left_indent = Inches(1.5)

        self.addNewLines(paragraph, amount = 5)

        # CONSOANTE DISPÕE
        paragraph = self.document.add_paragraph("Consoante dispõe a decisão exarada no evento __ dos autos supracitados, ")
        paragraph.add_run("in verbis: ").italic = True
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.5)

        self.addNewLines(paragraph, amount = 2)

        # COMPLEMENTO 1
        paragraph = self.document.add_paragraph(".,.,.,.,.,.,.,.,")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.left_indent = Inches(1.5)

        self.addNewLines(paragraph, amount = 1)

        # COMPLEMENTO 2
        paragraph = self.document.add_paragraph("    [...]")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.left_indent = Inches(1.5)

        self.addNewLines(paragraph, amount = 2)

        # ASSIM, CONSIDERANDO OS TERMOS
        paragraph = self.document.add_paragraph("Assim, considerando os termos da sentença proferida, tem-se que o valor devido perfaz o montante atualizado de ")
        paragraph.add_run("R$ ---.---,--- (----- ----- ----- ----)").underline = True
        paragraph.add_run(", sendo o montante atualizado de R$ __.___,__ referentes às restituições dos valores retidos indevidamente e R$ __.___,__ referente aos danos morais arbitrados, conforme demonstrativos anexos.")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.5)
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.addNewLines(paragraph, amount = 1)

        # ISTO POSTO
        paragraph = self.document.add_paragraph("Isto posto, requer o cumprimento da sentença na forma da legislação vigente, no que toca aos valores e cálculos acima citados, com a intimação da requerida para, querendo, no prazo de 30 (trinta) dias, impugnar a presente execução (art. 535, ")
        paragraph.add_run("caput").italic = True
        paragraph.add_run(", do CPC).")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.5)
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.addNewLines(paragraph, amount = 1)

        # TRANSCORRIDO
        paragraph = self.document.add_paragraph("Transcorrido o prazo acima assinalado sem impugnação ou rejeitadas as arguições da executada, requer seja expedido mandado dirigido à União, na pessoa de seu representante, para pagamento da importância de R$ __.___,__ (____ ___ ____ ____), acrescida de juros e correção monetária (art. 535, § 3º, II, do CPCP), mediante RPV, assinalando o prazo de até 60 (sessenta) dias para pagamento.")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.5)
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.addNewLines(paragraph, amount = 1)


        # NAO HAVENDO PAGAMENTO
        paragraph = self.document.add_paragraph("Não havendo o pagamento, requer a realização de penhora via Sisbajud dos ativos financeiros eventualmente existentes em nome da executada.")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.5)
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.addNewLines(paragraph, amount = 1)

        # PROTESTA
        paragraph = self.document.add_paragraph("Protesta, ainda, pela produção de todos os meios de prova em direito admitidos, em especial prova self.documental, pericial e testemunhal, cujo rol será oportunamente apresentado.")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.5)
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.addNewLines(paragraph, amount = 1)

        # POR FIM
        paragraph = self.document.add_paragraph("Por fim, pleiteia a condenação da requerida ao pagamento de eventuais custas e honorários advocatícios, eis que deu causa à presente.")
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.first_line_indent = Inches(0.5)
        paragraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.addNewLines(paragraph, amount = 4)

        # FOOTERS
        self.addFooters()

        # ESTILO
        self.setStyle()

        nome_documento = f"Cumprimento de Sentença {self.cliente_final}.docx"
        if len(nome_documento) > 42:
            nome_documento = nome_documento[:42] + ".docx"

        path_documento = os.path.join("arquivos/cumprimentos de sentenca", nome_documento)
        self.document.save(path_documento)
        os.startfile(path_documento)
        time.sleep(2)
        maximize = win32gui.GetForegroundWindow()
        win32gui.ShowWindow(maximize, win32con.SW_MAXIMIZE)
