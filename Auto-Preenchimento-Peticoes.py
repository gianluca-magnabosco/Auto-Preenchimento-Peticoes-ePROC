from concurrent.futures import process
from operator import index
import tkinter as tk
import klembord
from tkinter import *
from tkinter import messagebox
from tkinter.font import BOLD
from tkinter.ttk import *
from tkinter import ttk
from tkinter.constants import CENTER, TOP, DISABLED
from openpyxl import load_workbook
import pyexcel as p
import pyexcel_xls
import pyexcel_xlsx
import os
import glob
import winshell
from win32com.client import Dispatch
from functools import partial
import datetime
from datetime import date
import win32gui
import win32con
import time
import pandas as pd
import psycopg2
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

cliente_final = ''
local_path = os.getcwd()



# load file, sheet
wb = load_workbook('processos_final_ordenado.xlsx')
ws = wb.active


# row number variable
max_rows = ws.max_row

# append list variables
processos = []
cliente = []
adversa = []
cidade = []

# copy num processos
for i in range(1,max_rows+1):
    processos.append(ws.cell(row = i, column = 1).value)

# copy cliente
for i in range(1,max_rows+1):
    cliente.append(ws.cell(row = i, column = 2).value)

# copy parte adversa
for i in range(1,max_rows+1):
    adversa.append(ws.cell(row = i, column = 3).value)

# copy cidade
for i in range(1,max_rows+1):
    cidade.append(ws.cell(row = i, column = 4).value)


# append lists
numero_processo = []

# get numero_processo
for i in range(0,max_rows):
    numero_processo.append(processos[i])

numero_processo = [x[:25] for x in numero_processo]


diadehoje = date.today().strftime("%d/%m/%Y")

dia = diadehoje[:2]
mes = diadehoje[3:5]
ano = diadehoje[8:]

meses_extenso = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
mes = int(mes)
mes = meses_extenso[mes-1]


# -------------------------------------------------------------------------------------------------------------------------------

#### DEALING WITH ROOT SCREEN
# center window
def center_window(width=860,height=640):
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x = (screen_width/2) - (width/2)
    y = (screen_height/2) - (height/2)
    root.geometry('%dx%d+%d+%d' % (width, height, x, y))

root = Tk()
root.title("Auto Peticionamento")
center_window(860, 640)
bg = PhotoImage(file = "root_background.png")
background_label = tk.Label(root, image=bg,bg='white').place(relx=0.5,rely=0.5,anchor=CENTER)
root.resizable(False,False)
iconFile = 'icone.ico'
root.iconbitmap(default=iconFile)


# close program confirmation
def on_closeroot():
    close = messagebox.askokcancel("Confirmação", "Tem certeza que deseja fechar o programa?")
    if close:
        root.destroy()
root.protocol("WM_DELETE_WINDOW", on_closeroot)



# credits
feitopor = tk.Label(text="Programa criado por: Gianluca Notari Magnabosco da Silva",font=('',7),bg="white")
feitopor.pack()
feitopor.place(relx=0.84, rely=0.98, anchor=CENTER)



top_bg = PhotoImage(file = "top_background.png")
alterar_top_bg = PhotoImage(file = "alterar_top_background.png")
adicionar_top_bg = PhotoImage(file = "adicionar_top_background.png")

def insert_input():
    # MINIMIZE ROOT
    root.wm_state('iconic')
    # MAXIMIZE ROOT
    #root.wm_state('normal')

    #### DEALING WITH TOP (POP-UP) SCREEN
    # background image
    top = tk.Toplevel(root)
    global top_bg
    background_label = tk.Label(top, image=top_bg,bg='white').place(relx=0.5,rely=0.5,anchor=CENTER)

    # pop up close
    def on_close_top():
        top.destroy()
        root.wm_state('normal')
    top.protocol("WM_DELETE_WINDOW", on_close_top)

    # center pop up window
    def center_window_pop_up(width=360, height=150):
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        x = (screen_width/2) - (width/2)
        y = (screen_height/2) - (height/2)
        top.geometry('%dx%d+%d+%d' % (width, height, x, y))
    center_window_pop_up(360, 150)

    top.title("Gerador de Petição")
    #top.attributes("-topmost", True)
    #root.attributes("-topmost", False)
    top.resizable(False,False)
    tk.Label(top, text= "Insira o número do processo:",font=('Arial',9),bg='white').place(relx=0.5,rely=0.14,anchor=CENTER)

    def is_peticao(index):
        cidade_final = cidade[index]
        num_processo_final = numero_processo[index]

        #### PETIÇÃO INICIAL

        document = Document()

        # TITULO
        p = document.add_paragraph()
        p.add_run("EXMO. SR. DR. JUIZ DE DIREITO DA _ª VARA CÍVEL DA COMARCA DE {}/SC".format(cidade_final.upper())).bold = True

        p_format = p.paragraph_format

        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()


        # AUTOS
        p = document.add_paragraph()
        p.add_run("Autos nº: {}".format(num_processo_final)).bold = True
        p.paragraph_format.line_spacing_rule = 0


        # PARTE ADVERSA
        p = document.add_paragraph()
        p.add_run("Parte adversa: ").bold = True
        p.add_run("{}".format(adversa_final).title())

        run = p.add_run()

        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()




        # CLIENTE
        p = document.add_paragraph()
        p.add_run("{}, ".format(cliente_final)).bold = True
        p.add_run("vem respeitosamente à presença de Vossa Excelência através de seu procurador que esta subscreve, expor e requerer o que segue: ")


        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.5)
        p_format.left_indent = Inches(1.5)

        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.line_spacing_rule = 0

        run = p.add_run()

        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()


        # TERMOS
        p = document.add_paragraph("Nestes termos,")
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # DEFERIMENTO
        p = document.add_paragraph("Pede deferimento.")
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # DIA DE HOJE
        p = document.add_paragraph("Mafra, {} de {} de 20{}.". format(dia, mes, ano))
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        run.add_break()
        run.add_break()



        # FOOTER MARCIO
        p = document.add_paragraph()
        p.add_run("MÁRCIO MAGNABOSCO DA SILVA").bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # FOOTER OAB MARCIO
        p = document.add_paragraph()
        p.add_run("OAB/SC 9.738 – OAB/PR 20.962").bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        run.add_break()
        run.add_break()


        # FOOTER ALINE
        p = document.add_paragraph()
        p.add_run("ALINE REWAY RUTHES").bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # FOOTER OAB ALINE
        p = document.add_paragraph()
        p.add_run("OAB/SC 52.034").bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER



        # AJEITANDO ESTILO
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)


        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']
            paragraph.paragraph_format.space_after = Pt(0)
            if 'Nestes' in paragraph.text:
                paragraph.paragraph_format.space_after = Pt(1.5)
            if 'Pede' in paragraph.text:
                paragraph.paragraph_format.space_after = Pt(1.5)
            if 'Mafra' in paragraph.text:
                paragraph.paragraph_format.space_after = Pt(1.5)
        for run in paragraph.runs:
            run.font.size = Pt(12)



        nome_documento = "Petição {}.docx".format(cliente_final)
        if len(nome_documento) > 42:
            nome_documento = nome_documento[:42] + ".docx"

        document.save(nome_documento)
        documento = os.path.join(local_path, nome_documento)
        os.startfile(documento)
        time.sleep(2)
        maximize = win32gui.GetForegroundWindow()
        win32gui.ShowWindow(maximize, win32con.SW_MAXIMIZE)


    def is_sentenca(index):
        cidade_final = cidade[index]
        num_processo_final = numero_processo[index]



        #### CUMPRIMENTO DE SENTENÇA

        document = Document()


        # TITULO
        p = document.add_paragraph()
        p.add_run("EXMO. SR. JUIZ FEDERAL DA _ª VARA FEDERAL DE {} - SEÇÃO JUDICIÁRIA DE SANTA CATARINA".format(cidade_final.upper())).bold = True

        p_format = p.paragraph_format

        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()


        # AUTOS
        p = document.add_paragraph()
        p.add_run("Autos nº: {}".format(num_processo_final)).bold = True
        p.paragraph_format.line_spacing_rule = 0


        # CUMPR SENTENÇA
        p = document.add_paragraph()
        p.add_run("Cumprimento de Sentença").bold = True
        p.paragraph_format.line_spacing_rule = 0


        p = document.add_paragraph()
        p.add_run("_______________________________").bold = True
        p.paragraph_format.line_spacing_rule = 0

        run = p.add_run()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()

        # CLIENTE
        p = document.add_paragraph()
        p.add_run("{}, ".format(cliente_final)).bold = True
        p.add_run("____________, através dos procuradores que a esta subscrevem, MÁRCIO MAGNABOSCO DA SILVA, advogado inscrito na Ordem dos Advogados do Brasil sob nº 9.738/SC e 20.962/PR e ALINE REWAY RUTHES, advogada inscrita na Ordem dos Advogados do Brasil sob nº 52.034/SC, com escritório profissional na rua Felipe Schmidt, nº 354, conjunto nº 01, Mafra/SC, vem, respeitosamente perante Vossa Excelência, nos termos do art. 513 e seguintes do CPC, requerer")


        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.25)
        p_format.left_indent = Inches(1.5)

        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.line_spacing_rule = 0

        run = p.add_run()

        run.add_break()
        run.add_break()



        # CUMPR SENTENÇA
        p = document.add_paragraph()
        p.add_run("CUMPRIMENTO DE SENTENÇA").bold = True

        p_format = p.paragraph_format
        p_format.left_indent = Inches(1.5)
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        run.add_break()
        run.add_break()


        # EM FACE DE
        p = document.add_paragraph("em face de")

        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.1)
        p_format.left_indent = Inches(1.5)

        run = p.add_run()

        run.add_break()
        run.add_break()


        # PARTE ADVERSA
        p = document.add_paragraph()
        p.add_run("{}, ".format(adversa_final)).bold = True
        p.add_run("______ em razão dos fatos e fundamentos a seguir aduzidos: ")


        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.25)
        p_format.left_indent = Inches(1.5)

        run = p.add_run()

        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()


        # CONSOANTE DISPÕE
        p = document.add_paragraph("Consoante dispõe a decisão exarada no evento __ dos autos supracitados, ")
        p.add_run("in verbis: ").italic = True

        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.5)

        run = p.add_run()

        run.add_break()
        run.add_break()


        # COMPLEMENTO 1
        p = document.add_paragraph(".,.,.,.,.,.,.,.,")
        p_format = p.paragraph_format
        p_format.left_indent = Inches(1.5)

        run = p.add_run()

        run.add_break()


        # COMPLEMENTO 2
        p = document.add_paragraph("    [...]")
        p_format = p.paragraph_format
        p_format.left_indent = Inches(1.5)

        run = p.add_run()

        run.add_break()
        run.add_break()

        # ASSIM, CONSIDERANDO OS TERMOS
        p = document.add_paragraph("Assim, considerando os termos da sentença proferida, tem-se que o valor devido perfaz o montante atualizado de ")
        p.add_run("R$ ---.---,--- (----- ----- ----- ----)").underline = True
        p.add_run(", sendo o montante atualizado de R$ __.___,__ referentes às restituições dos valores retidos indevidamente e R$ __.___,__ referente aos danos morais arbitrados, conforme demonstrativos anexos.")

        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.5)
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run()

        run.add_break()


        # ISTO POSTO
        p = document.add_paragraph("Isto posto, requer o cumprimento da sentença na forma da legislação vigente, no que toca aos valores e cálculos acima citados, com a intimação da requerida para, querendo, no prazo de 30 (trinta) dias, impugnar a presente execução (art. 535, ")
        p.add_run("caput").italic = True
        p.add_run(", do CPC).")

        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.5)
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run()

        run.add_break()


        # TRANSCORRIDO
        p = document.add_paragraph("Transcorrido o prazo acima assinalado sem impugnação ou rejeitadas as arguições da executada, requer seja expedido mandado dirigido à União, na pessoa de seu representante, para pagamento da importância de R$ __.___,__ (____ ___ ____ ____), acrescida de juros e correção monetária (art. 535, § 3º, II, do CPCP), mediante RPV, assinalando o prazo de até 60 (sessenta) dias para pagamento.")

        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.5)
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run()

        run.add_break()


        # NAO HAVENDO PAGAMENTO
        p = document.add_paragraph("Não havendo o pagamento, requer a realização de penhora via Sisbajud dos ativos financeiros eventualmente existentes em nome da executada.")

        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.5)
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run()

        run.add_break()


        # PROTESTA
        p = document.add_paragraph("Protesta, ainda, pela produção de todos os meios de prova em direito admitidos, em especial prova documental, pericial e testemunhal, cujo rol será oportunamente apresentado.")

        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.5)
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run()

        run.add_break()


        # POR FIM
        p = document.add_paragraph("Por fim, pleiteia a condenação da requerida ao pagamento de eventuais custas e honorários advocatícios, eis que deu causa à presente.")

        p_format = p.paragraph_format
        p_format.first_line_indent = Inches(0.5)
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = p.add_run()

        run.add_break()
        run.add_break()
        run.add_break()
        run.add_break()


        # TERMOS
        p = document.add_paragraph("Nestes termos,")
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # DEFERIMENTO
        p = document.add_paragraph("Pede deferimento.")
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # DIA DE HOJE
        p = document.add_paragraph("Mafra, {} de {} de 20{}.". format(dia, mes, ano))
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        run.add_break()
        run.add_break()



        # FOOTER MARCIO
        p = document.add_paragraph()
        p.add_run("MÁRCIO MAGNABOSCO DA SILVA").bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # FOOTER OAB MARCIO
        p = document.add_paragraph()
        p.add_run("OAB/SC 9.738 – OAB/PR 20.962").bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        run.add_break()
        run.add_break()


        # FOOTER ALINE
        p = document.add_paragraph()
        p.add_run("ALINE REWAY RUTHES").bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # FOOTER OAB ALINE
        p = document.add_paragraph()
        p.add_run("OAB/SC 52.034").bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER



        # AJEITANDO ESTILO
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)


        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']
            paragraph.paragraph_format.space_after = Pt(0)
            if 'Nestes' in paragraph.text:
                paragraph.paragraph_format.space_after = Pt(1.5)
            if 'Pede' in paragraph.text:
                paragraph.paragraph_format.space_after = Pt(1.5)
            if 'Mafra' in paragraph.text:
                paragraph.paragraph_format.space_after = Pt(1.5)
        for run in paragraph.runs:
            run.font.size = Pt(12)


        nome_documento = "Cumprimento de Sentença {}.docx".format(cliente_final)
        if len(nome_documento) > 42:
            nome_documento = nome_documento[:42] + ".docx"

        document.save(nome_documento)
        documento = os.path.join(local_path, nome_documento)
        os.startfile(documento)
        time.sleep(2)
        maximize = win32gui.GetForegroundWindow()
        win32gui.ShowWindow(maximize, win32con.SW_MAXIMIZE)



    def validateInput(num_processo):
        peticao = it_is_peticao.get()
        sentenca = cumprimento_sentenca.get()
        processo_atual = num_processo.get()
        
        if len(processo_atual) == 20:
            edit_num_processo = list(str(processo_atual))
            edit_num_processo.insert(7, '-')
            edit_num_processo.insert(10, '.')
            edit_num_processo.insert(15, '.')
            edit_num_processo.insert(17, '.')
            edit_num_processo.insert(20, '.')
            processo_atual = "".join(edit_num_processo)
        if len(processo_atual) > 25:
            tamanho = len(processo_atual)
            edit_num_processo = list(str(processo_atual))
            for i in range(25,tamanho):
                edit_num_processo.pop()
            processo_atual = "".join(edit_num_processo)
        
            


        try:
            index = numero_processo.index(processo_atual)
        except:
            if (peticao != 0 or sentenca != 0) and peticao != sentenca:
                tk.messagebox.showerror(title="Erro", message="Processo não encontrado!") 
        else:
            if (peticao != 0 or sentenca != 0) and peticao != sentenca:
                cliente_temp = cliente[index]
                adversa_temp = adversa[index]

                
                # establish connection to postgres database
                pg = psycopg2.connect(
                host = "localhost",
                database = "Processos",
                user = "postgres",
                password = "")


                # set cursor
                query = pg.cursor()
                # sample queries
                query.execute("""SELECT num_processo FROM processos WHERE num_processo = '{}' """.format(processo_atual))
                # fetch all queries and show
                processo_db = query.fetchall()
                global cliente_final
                global adversa_final
                # if list does not have content (query returns false), do something
                if not processo_db:
                    is_cliente = tk.messagebox.askyesnocancel(title="Processo nº {}".format(processo_atual), message="{} é seu cliente nesse processo?".format(cliente_temp))
                    if is_cliente:
                        cliente_final = cliente_temp

                        adversa_final = adversa_temp

                    if not is_cliente:
                        cliente_final = adversa_temp

                        adversa_final = cliente_temp


                    query.execute("INSERT INTO processos (num_processo, cliente, parte_adversa, cidade) values ('{}', '{}', '{}', '{}')".format(numero_processo[index], cliente_final, adversa_final, cidade[index]))






                    
                    # commit changes to db
                    pg.commit()

                    # close cursor to db
                    query.close()


                if processo_db:
                    # set cursor
                    query = pg.cursor()
                    # sample queries
                    query.execute("""SELECT * FROM processos WHERE num_processo = '{}' """.format(processo_atual))
                    
                    list = query.fetchall()

                    cliente_final = list[0][1]
                    adversa_final = list[0][2]

                    print(list[0][1], list[0][2])

                    # commit changes to db
                    pg.commit()

                    # close cursor to db
                    query.close()

                
                
                # close connection to db
                pg.close()
                tk.messagebox.showinfo(title="Sucesso!", message="Arquivo gerado com sucesso!")




        if peticao > 0 and sentenca > 0:
            tk.messagebox.showwarning(title="Erro", message="Selecione uma opção por vez!")
        if peticao == 0 and sentenca == 0:
            tk.messagebox.showwarning(title="Erro", message="Selecione ao menos uma opção!") 
        if peticao > 0 and sentenca == 0:
            is_peticao(index)
        if sentenca > 0 and peticao == 0:
            is_sentenca(index)
            
        return processo_atual

    #num_processo text entry box
    num_processo = StringVar()
    num_processo_entry = Entry(top, textvariable=num_processo).place(relx=0.6,rely=0.36,anchor=CENTER,width=200)


    cumprimento_sentenca = IntVar()
    tk.Checkbutton(top, variable=cumprimento_sentenca,bg='#cde3f1').place(relx=0.1486,rely=0.54,anchor=CENTER)
    it_is_peticao = IntVar()
    tk.Checkbutton(top, variable=it_is_peticao,bg='#cde3f1').place(relx=0.7,rely=0.54,anchor=CENTER)

    validateInput = partial(validateInput, num_processo)        

    #num_processo button
    st = ttk.Style()
    st.configure('W.TButton', background='#a3cae7', foreground='black', font=('Open Sans',9))
    confirm_button = ttk.Button(top,style='W.TButton', text="Confirma", command=validateInput).place(relx=0.5,rely=0.79,anchor=CENTER,width=60) 
        
    # add RETURN key handler
    def handler(e):
        validateInput()
    top.bind('<Return>', handler)
    
  



# abrir pasta de petições
def open_directory():
    os.startfile(local_path)
    time.sleep(2)
    maximize = win32gui.GetForegroundWindow()
    win32gui.ShowWindow(maximize, win32con.SW_MAXIMIZE)



# alterar cliente
def alterar_cliente():
    root.wm_state('iconic')

    def validateInput2(num_processo2):
        alterar_top.attributes("-topmost", False)
        processo_atual = num_processo2.get()
        
        if len(processo_atual) == 20:
            edit_num_processo = list(str(processo_atual))
            edit_num_processo.insert(7, '-')
            edit_num_processo.insert(10, '.')
            edit_num_processo.insert(15, '.')
            edit_num_processo.insert(17, '.')
            edit_num_processo.insert(20, '.')
            processo_atual = "".join(edit_num_processo)
        if len(processo_atual) > 25:
            tamanho = len(processo_atual)
            edit_num_processo = list(str(processo_atual))
            for i in range(25,tamanho):
                edit_num_processo.pop()
            processo_atual = "".join(edit_num_processo)
        
            

        # establish connection to postgres database
        pg = psycopg2.connect(
        host = "localhost",
        database = "Processos",
        user = "postgres",
        password = "")


        # set cursor
        query = pg.cursor()
        
        try:
            query.execute("""SELECT num_processo FROM processos WHERE num_processo = '{}' """.format(processo_atual))
            exists = query.fetchall()
            index = numero_processo.index(processo_atual)
            
            #teste = exists[0][0]

        except:
            tk.messagebox.showerror(title="Erro", message="Processo não encontrado!") 
            alterar_top.attributes("-topmost", True)
        else:

            try:
                teste = exists[0][0]
            except:
                tk.messagebox.showwarning(title="Erro", message="Processo não consta no banco de dados!\nGere um arquivo com este número de processo para adicioná-lo ao banco de dados!")
                alterar_top.attributes("-topmost", True)
            else:
                # sample queries
                query.execute("""SELECT * FROM processos WHERE num_processo = '{}' """.format(processo_atual))
            
                # fetch all queries and assign
                processo_db = query.fetchall()
                cliente_atual = processo_db[0][1]
                parte_adversa_atual = processo_db[0][2]
                
                
                if processo_db:
                    is_cliente = tk.messagebox.askyesno(title="Processo nº {}".format(processo_atual), message="{} é seu cliente nesse processo?".format(cliente_atual))
                    
                    if not is_cliente:
                        query.execute("UPDATE processos SET cliente = '{}' WHERE num_processo = '{}'".format(parte_adversa_atual, processo_atual))
                        query.execute("UPDATE processos SET parte_adversa = '{}' WHERE num_processo = '{}'".format(cliente_atual, processo_atual))
                        tk.messagebox.showinfo(title="Sucesso!", message="Processo atualizado com sucesso!")
                        
                    # commit changes to db
                    pg.commit()

                    # close cursor to db
                    query.close()
                    
                    # close connection to db
                    pg.close()

                    if is_cliente:
                        tk.messagebox.showerror(title="Nada foi alterado", message="Nada foi atualizado!")

                on_close_alterar_top()


    alterar_top = tk.Toplevel(root)
    global alterar_top_bg
    background_label = tk.Label(alterar_top, image=alterar_top_bg,bg='white').place(relx=0.5,rely=0.5,anchor=CENTER)

    # pop up close
    def on_close_alterar_top():
        alterar_top.destroy()
        root.attributes("-topmost", True)
        root.wm_state('normal')
    alterar_top.protocol("WM_DELETE_WINDOW", on_close_alterar_top)

    # center pop up window
    def center_window_pop_up(width=360, height=150):
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        x = (screen_width/2) - (width/2)
        y = (screen_height/2) - (height/2)
        alterar_top.geometry('%dx%d+%d+%d' % (width, height, x, y))
    center_window_pop_up(340, 115)

    alterar_top.title("Alterar cliente")
    #top.attributes("-topmost", True)
    #root.attributes("-topmost", False)
    alterar_top.resizable(False,False)
    tk.Label(alterar_top, text= "Insira o número do processo que deseja alterar:",font=('Arial',9),bg='white').place(relx=0.5,rely=0.14,anchor=CENTER)

    num_processo2 = StringVar()
    num_processo2_entry = Entry(alterar_top, textvariable=num_processo2).place(relx=0.66,rely=0.49,anchor=CENTER,width=200)

    validateInput2 = partial(validateInput2, num_processo2) 
    
    #num_processo button
    st = ttk.Style()
    st.configure('W.TButton', background='#a3cae7', foreground='black', font=('Open Sans',9))
    confirm_button = ttk.Button(alterar_top,style='W.TButton', text="Confirma", command=validateInput2).place(relx=0.5,rely=0.79,anchor=CENTER,width=60) 
        

   
    # add RETURN key handler
    def handler(e):
        validateInput2()
    alterar_top.bind('<Return>', handler)



def adicionar_processo():
    root.wm_state('iconic')

    def validateInput3(num_processo2, cliente_input, adversa_input, cidade_input):
        adicionar_top.attributes("-topmost", False)
        
        processo_atual = num_processo2.get()
        cliente = cliente_input.get()
        adversa = adversa_input.get()
        cidade = cidade_input.get()
        
        if len(processo_atual) == 20:
            edit_num_processo = list(str(processo_atual))
            edit_num_processo.insert(7, '-')
            edit_num_processo.insert(10, '.')
            edit_num_processo.insert(15, '.')
            edit_num_processo.insert(17, '.')
            edit_num_processo.insert(20, '.')
            processo_atual = "".join(edit_num_processo)
        if len(processo_atual) > 25:
            tamanho = len(processo_atual)
            edit_num_processo = list(str(processo_atual))
            for i in range(25,tamanho):
                edit_num_processo.pop()
            processo_atual = "".join(edit_num_processo)
        
            

        # establish connection to postgres database
        pg = psycopg2.connect(
        host = "localhost",
        database = "Processos",
        user = "postgres",
        password = "")


        # set cursor
        query = pg.cursor()
        
        try: 
            query.execute("INSERT INTO processos(num_processo, cliente, parte_adversa, cidade) VALUES ('{}', '{}', '{}', '{}')".format(processo_atual, cliente, adversa, cidade))
        except:
            tk.messagebox.showwarning(title="Erro", message="Erro!\nTente novamente")
            adicionar_top.attributes("-topmost", True)           
        else: 
            tk.messagebox.showinfo(title="Sucesso!", message="Processo adicionado com sucesso!")                

                

                        
        # commit changes to db
        pg.commit()

        # close cursor to db
        query.close()
        
        # close connection to db
        pg.close()

        on_close_adicionar_top()


    adicionar_top = tk.Toplevel(root)
    global adicionar_top_bg
    background_label = tk.Label(adicionar_top, image=adicionar_top_bg,bg='white').place(relx=0.5,rely=0.5,anchor=CENTER)

    # pop up close
    def on_close_adicionar_top():
        adicionar_top.destroy()
        root.attributes("-topmost", True)
        root.wm_state('normal')
    adicionar_top.protocol("WM_DELETE_WINDOW", on_close_adicionar_top)

    # center pop up window
    def center_window_pop_up(width, height):
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        x = (screen_width/2) - (width/2)
        y = (screen_height/2) - (height/2)
        adicionar_top.geometry('%dx%d+%d+%d' % (width, height, x, y))
    center_window_pop_up(520, 250)

    adicionar_top.title("Adicionar processo")
    #top.attributes("-topmost", True)
    #root.attributes("-topmost", False)
    adicionar_top.resizable(False,False)
    tk.Label(adicionar_top, text= "Insira os dados referentes ao processo que deseja inserir:",font=('Arial',12, BOLD),bg='white').place(relx=0.5,rely=0.14,anchor=CENTER)

    num_processo2 = StringVar()
    num_processo2_entry = Entry(adicionar_top, textvariable=num_processo2).place(relx=0.54,rely=0.33,anchor=CENTER,width=251)
    
    cliente_input = StringVar()
    cliente_input_entry = Entry(adicionar_top, textvariable = cliente_input).place(relx=0.51, rely=0.48, anchor=CENTER, width=285)

    adversa_input = StringVar()
    adversa_input_entry = Entry(adicionar_top, textvariable = adversa_input).place(relx=0.52, rely=0.63, anchor=CENTER, width=277)

    cidade_input = StringVar()
    cidade_input_entry = Entry(adicionar_top, textvariable = cidade_input).place(relx=0.51, rely=0.79, anchor=CENTER, width=285)


    validateInput3 = partial(validateInput3, num_processo2, cliente_input, adversa_input, cidade_input) 
    
    #num_processo button
    st = ttk.Style()
    st.configure('W.TButton', background='#a3cae7', foreground='black', font=('Open Sans',9))
    confirm_button = ttk.Button(adicionar_top, style='W.TButton', text="Confirma", command=validateInput3).place(relx=0.5,rely=0.93,anchor=CENTER,width=60) 
        

   
    # add RETURN key handler
    def handler(e):
        validateInput3()
    adicionar_top.bind('<Return>', handler)


    

# open directory button
st2 = Style()
st2.configure('B.TButton', background='white', foreground='black', font=('Arial', 9))
button1 = Button(root, style='B.TButton', text='Abrir',command=open_directory,width=27.75)
button1.pack()
button1.place(relx=0.27, rely=0.85, anchor=CENTER)

# modify client button
st3 = Style()
st3.configure('C.TButton', background='white', foreground='black', font=('Helvetica', 9))
button2 = Button(root, style='C.TButton', text='Alterar',command=alterar_cliente,width=27.75)
button2.pack()
button2.place(relx=0.83, rely=0.44, anchor=CENTER)

# order excel button
st4 = Style()
st4.configure('D.TButton', background='white', foreground='black', font=('Helvetica', 9))
button3 = Button(root, style='D.TButton', text='Adicionar',command=adicionar_processo,width=27.75)
button3.pack()
button3.place(relx=0.81, rely=0.73, anchor=CENTER)



# run code button
st5 = Style()
st5.configure('E.TButton', background='white', foreground='black', font=('Helvetica', 12))
button4 = Button(root, style='E.TButton', text='Gerar Petições',command=insert_input,width=27.75)
button4.pack()
button4.place(relx=0.46, rely=0.665, anchor=CENTER)


insert_input()


root.mainloop()